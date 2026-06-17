from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Development_of_Flutter_Apps_for_Fourier_Transform_Dissertation_v5_text_diagrams.docx"


def shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def table_width(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")

    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for w in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[i]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        shade(cell, "E8EEF5")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value
    table_width(table, widths)
    doc.add_paragraph()


def add_boxed_flow(doc: Document, caption: str, steps: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(steps) * 2 - 1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, (title, body) in enumerate(steps):
        row_idx = i * 2
        cell = table.rows[row_idx].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = True
        p.add_run("\n" + body)
        shade(cell, "E8EEF5" if i % 2 == 0 else "F4F6F9")
        cell_margins(cell, top=120, bottom=120, start=160, end=160)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        if i < len(steps) - 1:
            arrow_cell = table.rows[row_idx + 1].cells[0]
            arrow_cell.text = "v"
            arrow_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell_margins(arrow_cell, top=20, bottom=20, start=120, end=120)
    table_width(table, [7200])
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(80, 90, 100)
    doc.add_paragraph()


def add_two_path_diagram(doc: Document, caption: str) -> None:
    table = doc.add_table(rows=6, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ["Symbolic Path (Backend)", "", "Numerical Path (Flutter Local)"],
        ["User expression", "same input", "User expression"],
        ["POST /fourier", "separate processing", "Local parser and evaluator"],
        ["Rule-first backend", "complementary", "Sample signal on time grid"],
        ["result_latex + steps_latex", "not competing", "FFT / Riemann-sum view"],
        ["Flutter symbolic render", "learning app combines both", "Interactive 4+1 charts"],
    ]
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell, top=100, bottom=100, start=100, end=100)
            if r_idx == 0:
                shade(cell, "E8EEF5" if c_idx != 1 else "FFFFFF")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            elif c_idx == 1:
                shade(cell, "FFFFFF")
            else:
                shade(cell, "F4F6F9")
    table_width(table, [3900, 1560, 3900])
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(80, 90, 100)
    doc.add_paragraph()


def add_hybrid_boundary_diagram(doc: Document, caption: str) -> None:
    table = doc.add_table(rows=5, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    rows = [
        ["Flutter Application", "API Boundary", "Python FastAPI Backend"],
        ["User interaction\nkeypad, examples, navigation", "symbolic request ->", "POST /fourier endpoint\nexpression parsing"],
        ["Numerical visualisation\nsignal, kernel, integrand, scan", "", "Rule-first symbolic engine\npairs, properties, distributions"],
        ["Symbolic display\nx(t), X(omega), steps", "<- LaTeX response", "Response formatter\nresult_latex, steps_latex"],
        ["Main smart-device learning interface", "clear separation", "Supporting symbolic service"],
    ]
    for r_idx, row in enumerate(rows):
        for c_idx, text in enumerate(row):
            cell = table.rows[r_idx].cells[c_idx]
            cell.text = text
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell_margins(cell, top=100, bottom=100, start=100, end=100)
            if r_idx == 0:
                shade(cell, "E8EEF5" if c_idx == 0 else "DDEFE4" if c_idx == 2 else "FFFFFF")
                for run in cell.paragraphs[0].runs:
                    run.bold = True
            elif c_idx == 1:
                shade(cell, "FFFFFF")
            else:
                shade(cell, "F4F6F9")
    table_width(table, [3900, 1560, 3900])
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(80, 90, 100)
    doc.add_paragraph()


def add_chart_layout_diagram(doc: Document, caption: str) -> None:
    table = doc.add_table(rows=3, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    top = [
        "1. Original Signal\nx(t)",
        "2. Kernel\nexp(-j omega t)",
        "3. Integrand\nx(t)exp(-j omega t)",
        "4. Running Integral\ncumulative sum",
    ]
    for i, text in enumerate(top):
        cell = table.rows[0].cells[i]
        cell.text = text
        shade(cell, "E8EEF5")
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_margins(cell, top=120, bottom=120, start=80, end=80)
    for i in range(4):
        cell = table.rows[1].cells[i]
        cell.text = "->" if i < 3 else "v sweep omega"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        cell_margins(cell, top=40, bottom=40, start=80, end=80)
    merged = table.rows[2].cells[0]
    for i in range(1, 4):
        merged = merged.merge(table.rows[2].cells[i])
    merged.text = "5. Frequency Scan: X(omega) over selected omega range\nThe omega slider links symbolic formulas with numerical intuition."
    shade(merged, "DDEFE4")
    merged.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    merged.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cell_margins(merged, top=140, bottom=140, start=100, end=100)
    table_width(table, [2340, 2340, 2340, 2340])
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9.5)
    r.font.color.rgb = RGBColor(80, 90, 100)
    doc.add_paragraph()


def add_code(doc: Document, code: str) -> None:
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph(style="CodeBlock")
        r = p.add_run(line.rstrip())
        r.font.name = "Consolas"
        r._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_note(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph(style="Callout")
    r = p.add_run(label + ": ")
    r.bold = True
    p.add_run(text)


def styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 14, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 10, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        s = doc.styles[name]
        s.font.name = "Calibri"
        s._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        s.font.bold = True
        s.font.size = Pt(size)
        s.font.color.rgb = color
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("CodeBlock", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)

    callout = doc.styles.add_style("Callout", 1)
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(6)


def title(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(110)
    r = p.add_run("Development of Flutter Apps for Fourier Transform")
    r.bold = True
    r.font.size = Pt(25)
    r.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("With Symbolic Derivation and Numerical Visualisation: Supported by a Rule-First Python Backend")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(46, 116, 181)

    for line in [
        "Dissertation Draft v5",
        "Restructured to keep the Flutter application as the main project focus",
        "Flutter frontend | Supporting Python FastAPI backend | Educational Fourier transform visualisation",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)
    doc.add_page_break()


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def numbers(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Number")


def build() -> None:
    doc = Document()
    styles(doc)
    title(doc)

    doc.add_heading("Abstract", level=1)
    para(
        doc,
        "This dissertation presents the development of a Flutter application for Fourier transform learning on smart devices. "
        "The original project requirement focuses on Flutter programming, simulator or device testing, step-by-step workings, and graphical illustrations of relevant Fourier transform equations. "
        "The project therefore keeps the Flutter application as the main user-facing system and treats the Python backend as a supporting symbolic computation layer.",
    )
    para(
        doc,
        "An early design option was to keep both interaction and symbolic processing inside Flutter/Dart. "
        "This approach was practical for input handling, numerical experimentation, and graph-based visualisation, but it was not suitable for reliable symbolic Fourier transform derivation involving distribution-theory objects such as Dirac delta, Heaviside step functions, principal value terms, and engineering-style LaTeX steps. "
        "The final system therefore adopts a hybrid architecture: Flutter provides the learning interface and numerical visualisation, while a rule-first Python FastAPI backend generates symbolic results and derivation steps.",
    )
    para(
        doc,
        "The current implementation includes structured examples, LaTeX formula display, a symbolic result page, numerical FT visualisation, teaching-oriented derivation steps, responsive local scrolling for long formulas, and automated tests. "
        "Current validation results are 68 passing backend tests and 6 passing Flutter tests, with 1 optional frontend-backend e2e test intentionally skipped unless explicitly enabled.",
    )

    doc.add_heading("Chapter 1: Introduction", level=1)
    doc.add_heading("1.1 Project Requirement and Motivation", level=2)
    para(
        doc,
        "The project requirement is to develop Flutter apps running on smart devices for Fourier transform. "
        "The application should be tested on a simulator or physical device and should show step-by-step workings to clarify Fourier transform solutions. "
        "Graph plots of relevant equations or illustrative graphics are also expected wherever possible.",
    )
    para(
        doc,
        "This requirement naturally positions the project as a learning application rather than a general-purpose symbolic calculator. "
        "The main design task is therefore to create a mobile-friendly interface that helps students input signals, inspect transform results, understand derivation steps, and connect symbolic formulas with visual behaviour.",
    )

    doc.add_heading("1.2 Problem Statement", level=2)
    para(
        doc,
        "Students learning Fourier transform often need more than a final answer. "
        "They need to see the transform pair, property, or distribution-theory argument that leads to the result. "
        "General-purpose computer algebra systems may return mathematically correct but pedagogically unclear expressions containing Piecewise branches, RootSum forms, complex conditions, or assumption-dependent simplifications.",
    )
    add_code(doc, r"Textbook form: F{1/(t^2+1)} = pi*exp(-abs(omega))")
    para(
        doc,
        "The project addresses this gap by combining Flutter-based learning interaction with a controlled symbolic computation layer that prioritises textbook-like engineering output.",
    )

    doc.add_heading("1.3 Aim and Objectives", level=2)
    para(
        doc,
        "The aim is to develop a Flutter-based Fourier transform learning app with symbolic derivation and numerical visualisation, supported by a rule-first Python backend where local Flutter symbolic processing is insufficient.",
    )
    bullets(
        doc,
        [
            "Develop a Flutter frontend for expression input, examples, LaTeX display, and graph-based learning.",
            "Provide step-by-step workings for supported Fourier transform questions.",
            "Evaluate why full symbolic derivation is not practical as a Flutter-only local feature.",
            "Use a Python backend to support distribution-aware symbolic derivation.",
            "Generate engineering-style result_latex and teaching-oriented steps_latex.",
            "Validate backend mathematical output and frontend layout behaviour with automated tests.",
        ],
    )

    doc.add_heading("1.4 Contributions", level=2)
    add_table(
        doc,
        ["Contribution", "Description"],
        [
            ["Flutter learning app", "Provides the main smart-device-oriented interface for input, examples, formulas, and visualisation."],
            ["Numerical visualisation", "Shows original signal, kernel, integrand, running integral, and frequency scan to support learning."],
            ["Symbolic derivation support", "Uses a Python backend to generate controlled Fourier transform results and derivation steps."],
            ["Engineering notation", "Displays j, u(t), PV, sign, compact absolute values, and delta derivatives in readable forms."],
            ["Small-screen handling", "Uses local scroll containers to prevent long formulas and notes from breaking narrow layouts."],
            ["Automated validation", "Tests mathematical results, teaching steps, formatting conventions, and frontend layout behaviour."],
        ],
        [2700, 6660],
    )

    doc.add_heading("Chapter 2: Background and Related Work", level=1)
    doc.add_heading("2.1 Fourier Transform Convention", level=2)
    add_code(doc, r"X(omega) = integral x(t) * exp(-j*omega*t) dt")
    para(
        doc,
        "The project uses the engineering Fourier transform convention with real angular frequency omega and imaginary unit j. "
        "This convention is also used by the frontend display layer and backend result formatter.",
    )

    doc.add_heading("2.2 Educational Difficulty", level=2)
    para(
        doc,
        "A common teaching difficulty is that the same Fourier transform can be represented in multiple equivalent symbolic forms. "
        "For students, the most useful form is usually the one that matches transform tables and classroom derivations. "
        "Therefore, the application emphasises readable transform pairs, properties, and derivation steps rather than raw symbolic integration output.",
    )

    doc.add_heading("2.3 Flutter/Dart Local Processing Limitation", level=2)
    para(
        doc,
        "Flutter and Dart are effective for building interactive interfaces, evaluating simple numerical expressions, and drawing plots. "
        "However, reliable symbolic Fourier transform derivation requires a computer algebra layer, rule matching, distribution objects, and controlled LaTeX formatting. "
        "Implementing these features entirely in Dart would require a substantial custom symbolic engine and a large validation effort.",
    )
    add_note(
        doc,
        "Design decision",
        "Flutter remains the main application layer, while Python is used only as a supporting service for symbolic derivation.",
    )

    doc.add_heading("Chapter 3: System Design", level=1)
    doc.add_heading("3.1 Hybrid Architecture", level=2)
    add_hybrid_boundary_diagram(
        doc,
        "Figure 3.1 Text-box architecture diagram: Flutter remains the main learning app, while Python supports symbolic derivation.",
    )
    add_code(
        doc,
        r"""
Flutter application
  -> keypad input, examples, LaTeX preview, charts
  -> POST /fourier
Python FastAPI backend
  -> parse expression
  -> rule-first transform derivation
  -> result_latex and steps_latex
Flutter application
  -> display x(t), X(omega), and derivation steps
""",
    )

    doc.add_heading("3.2 Frontend/Backend Boundary", level=2)
    add_two_path_diagram(
        doc,
        "Figure 3.2 Text-box diagram of the two complementary computation paths.",
    )
    add_table(
        doc,
        ["Layer", "Responsibility"],
        [
            ["Flutter frontend", "User input, examples, formula preview, symbolic result display, local numerical visualisation, responsive layout."],
            ["Python backend", "Expression parsing, rule-first symbolic transform, distribution rules, teaching steps, engineering LaTeX formatting."],
            ["Deployment boundary", "Frontend can be built in GitHub Actions; backend can be deployed separately on Render."],
        ],
        [2400, 6960],
    )

    doc.add_heading("3.3 Data Flow", level=2)
    numbers(
        doc,
        [
            "The user enters or selects a time-domain expression in the Flutter app.",
            "The frontend displays a LaTeX preview and can run local numerical visualisation.",
            "For symbolic derivation, the frontend sends the expression to the backend /fourier endpoint.",
            "The backend returns input_latex, result_latex, and steps_latex.",
            "The frontend renders the original function, transform result, and derivation steps.",
        ],
    )

    doc.add_heading("Chapter 4: Flutter Application Design", level=1)
    doc.add_heading("4.1 Home Page and Input Design", level=2)
    para(
        doc,
        "The home page is the main interaction point. "
        "It includes expression preview, keypad input, example navigation, FFT size controls, and access to symbolic and numerical result views. "
        "The keypad helps users enter expressions using supported syntax such as u(t), delta(t), sin, cos, exp, powers, and fractions.",
    )

    doc.add_heading("4.2 Structured Examples", level=2)
    para(
        doc,
        "The example list was expanded from simple expressions to a structured set of teaching examples. "
        "Each example includes an expression, title, category, input LaTeX, transform object, and description. "
        "This helps users understand the supported input range and reduces the chance of entering unsuitable expressions.",
    )
    add_code(
        doc,
        r"""
class FourierExample {
  final String expression;
  final String title;
  final String category;
  final String inputLatex;
  final String transformLatex;
  final String description;
}
""",
    )

    doc.add_heading("4.3 LaTeX Display and Responsive Layout", level=2)
    para(
        doc,
        "The frontend uses LaTeX rendering for formula preview, symbolic results, derivation steps, and general Fourier formulas. "
        "Long formulas and annotations can exceed the width of small screens, so local scroll containers were added. "
        "This avoids stretching the entire page and keeps the affected formula or note independently scrollable.",
    )
    add_code(
        doc,
        r"""
class ScrollableMathLine extends StatefulWidget {
  final String latex;
  final TextStyle? textStyle;
  final String? semanticsLabel;
}

class ScrollableTextLine extends StatefulWidget {
  final String text;
  final TextStyle? style;
  final String? semanticsLabel;
}
""",
    )

    doc.add_heading("4.4 Numerical Visualisation and FT Steps", level=2)
    add_chart_layout_diagram(
        doc,
        "Figure 4.1 Text-box layout of the 4+1 teaching visualisation page.",
    )
    para(
        doc,
        "The numerical visualisation page supports the original requirement for graph plots or illustrative graphics. "
        "It shows the signal, Fourier kernel, integrand, running integral, and frequency scan curve. "
        "This complements the symbolic backend by helping students see how the transform process behaves numerically.",
    )
    numbers(
        doc,
        [
            "Original signal x(t).",
            "Kernel e^{-j omega t}.",
            "Integrand x(t)e^{-j omega t}.",
            "Running integral or cumulative sum.",
            "Frequency scan curve.",
        ],
    )

    doc.add_heading("Chapter 5: Supporting Symbolic Backend", level=1)
    doc.add_heading("5.1 API Contract", level=2)
    add_boxed_flow(
        doc,
        "Figure 5.1 Text-box symbolic derivation flow from user input to Flutter rendering.",
        [
            ("1. User Input", "Expression selected or entered in the Flutter app."),
            ("2. API Request", "Flutter sends POST /fourier with the expression."),
            ("3. Parse and Normalise", "Backend converts syntax such as u(t), delta(t), and frac(a,b)."),
            ("4. Rule-First Derivation", "Known pairs, properties, and distribution rules are tried in priority order."),
            ("5. Format Output", "Backend returns engineering result_latex and teaching steps_latex."),
            ("6. Flutter Render", "The app displays x(t), X(omega), and derivation steps."),
        ],
    )
    add_code(
        doc,
        r"""
class FourierResponse(BaseModel):
    ok: bool
    input_latex: str
    result_latex: str
    steps_latex: list[str]
    error: str | None = None
    method: str | None = None
    form: str | None = None
""",
    )

    doc.add_heading("5.2 Rule-First Symbolic Method", level=2)
    add_boxed_flow(
        doc,
        "Figure 5.2 Text-box priority pipeline for rule-first symbolic derivation.",
        [
            ("1. Exact Matches", "Constants, Dirac delta, and unit step."),
            ("2. Known Transform Pairs", "Sin, cos, exponential, and rational pairs."),
            ("3. Distribution Rules", "PV reciprocal, t^n u(t), and delta derivatives."),
            ("4. Property-Based Derivation", "Linearity, shift, scale, and modulation."),
            ("5. Decomposition Rules", "Partial fractions, finite windows, and combinations."),
            ("6. Controlled Fallback", "Integral form or unsupported response."),
        ],
    )
    para(
        doc,
        "The backend does not use direct symbolic integration as the primary path. "
        "Instead, it matches known Fourier transform pairs and properties first. "
        "This allows the output to follow engineering notation and allows the backend to generate steps that explain the selected transform rule.",
    )
    bullets(
        doc,
        [
            "Dirac delta and shifted delta.",
            "Heaviside step and shifted step.",
            "Principal value reciprocal forms.",
            "Sinusoidal and exponential functions.",
            "Polynomial and polynomial-step distributions.",
            "Selected rational forms and finite windows.",
            "Linear combinations of supported pieces.",
        ],
    )

    doc.add_heading("5.3 Distribution-Aware Rule Example", level=2)
    add_code(
        doc,
        r"""
def _rule_poly_times_step_distribution(f):
    if not (isinstance(f, Mul) and f.has(Heaviside)):
        return None

    g = expand(simplify(f / h_t))
    if not g.is_polynomial(t):
        return None

    def _basis(n):
        if n == 0:
            return pi*DiracDelta(omega) - I*(1/omega)
        return (
            pi*(I**n)*Derivative(DiracDelta(omega), (omega, n))
            + ((-1)**n)*(I**(n-1))*factorial(n)*PV(1/(omega**(n+1)))
        )
""",
    )
    para(
        doc,
        "This rule demonstrates why the backend is useful: it handles distribution-style terms that would be difficult to implement robustly as a small Flutter-only symbolic engine.",
    )

    doc.add_heading("5.4 Teaching Step Generation", level=2)
    add_code(
        doc,
        r"""
def _teaching_steps(f, X, steps):
    cleaned = []
    for raw in steps or []:
        s = _format_step_display_latex(str(raw).strip())
        if "Method:" in s or "_rule_" in s or "matcher" in s:
            continue
        if "Piecewise" in s or "RootSum" in s or "meijerg" in s:
            continue
        cleaned.append(s)

    if r"X(\omega)=\int" not in "\n".join(cleaned):
        cleaned = _step_start_definition(f) + cleaned
    return cleaned
""",
    )
    para(
        doc,
        "The step generator removes internal implementation details and keeps the displayed derivation suitable for teaching. "
        "Recent formatting improvements also ensure that j, u(t), PV, sign, absolute values, and delta derivatives are displayed consistently.",
    )

    doc.add_heading("Chapter 6: Implementation Highlights", level=1)
    doc.add_heading("6.1 Frontend to Backend Request", level=2)
    add_code(
        doc,
        r"""
final res = await http.post(
  uri,
  headers: {'Content-Type': 'application/json'},
  body: jsonEncode({'expression': expression}),
);
""",
    )
    para(
        doc,
        "The frontend obtains symbolic derivations by sending the expression to the backend. "
        "The backend URL can be configured with API_BASE_URL, which supports separate frontend and backend deployment.",
    )

    doc.add_heading("6.2 Code Selection Rationale", level=2)
    add_table(
        doc,
        ["Code excerpt", "Why it is included"],
        [
            ["FourierExample", "Shows how preset examples became structured teaching aids."],
            ["ScrollableMathLine", "Shows the recent small-screen formula overflow solution."],
            ["FourierResponse", "Shows the result/steps contract between backend and frontend."],
            ["_rule_poly_times_step_distribution", "Shows a representative distribution-aware symbolic rule."],
            ["_teaching_steps", "Shows how backend output is converted into teaching-friendly derivation steps."],
        ],
        [3300, 6060],
    )

    doc.add_heading("Chapter 7: Testing and Evaluation", level=1)
    doc.add_heading("7.1 Backend Tests", level=2)
    add_code(
        doc,
        r"""
cd backend
.\.venv\Scripts\python.exe -m pytest tests -q

Result: 68 passed
""",
    )
    bullets(
        doc,
        [
            "Strict result_latex checks for implemented transform pairs.",
            "Teaching-style steps_latex checks.",
            "Step snapshot tests for representative examples.",
            "Guards against Piecewise, RootSum, arg, polar_lift, meijerg, matcher, debug, and srepr.",
            "Engineering notation checks for j, u(t), PV, sign, absolute values, and delta derivatives.",
        ],
    )

    doc.add_heading("7.2 Flutter Tests", level=2)
    add_code(
        doc,
        r"""
cd flutter_app
flutter test

Result: 6 passed, 1 skipped
""",
    )
    bullets(
        doc,
        [
            "Responsive breakpoint checks.",
            "Home page rendering and example navigation.",
            "Multiple screen sizes without layout exceptions.",
            "Local scroll containers for long formulas and notes.",
            "FT Steps long-formula display on narrow screens.",
            "Optional frontend-backend e2e test skipped unless explicitly enabled.",
        ],
    )

    doc.add_heading("7.3 Evaluation Limitations", level=2)
    para(
        doc,
        "The tests validate implemented rule families and known frontend behaviours. "
        "They do not prove correctness for every possible user expression. "
        "The optional e2e test is separated from normal frontend testing so that GitHub Actions frontend builds are not dependent on the Render backend being live.",
    )

    doc.add_heading("Chapter 8: Discussion", level=1)
    doc.add_heading("8.1 Supervisor Concerns Addressed", level=2)
    add_table(
        doc,
        ["Concern", "How it is addressed"],
        [
            ["General user input", "The app accepts user expressions and matches general transform families within a defined engineering range."],
            ["Generic Fourier transforms", "The backend uses parameterized rules rather than only fixed examples."],
            ["Step-by-step workings", "steps_latex is generated and cleaned for teaching use."],
            ["Graph plots", "The FT Steps page provides signal, kernel, integrand, running integral, and frequency scan plots."],
            ["Flutter project focus", "Flutter remains the main application; Python is explained as a supporting symbolic service."],
            ["Clear response file", "Detailed and brief Q&A documents are maintained in docs/QA.md and docs/QA_email_brief.md."],
        ],
        [3000, 6360],
    )

    doc.add_heading("8.2 Capability and Scope", level=2)
    para(
        doc,
        "The application is not an unrestricted CAS. "
        "It supports a defined engineering range including constants, impulses, steps, finite windows, trigonometric functions, exponentials, PV reciprocal forms, polynomial distributions, polynomial-step forms, selected rational forms, and combinations of supported terms. "
        "Unsupported expressions should return a controlled fallback or error rather than an unreliable symbolic claim.",
    )

    doc.add_heading("8.3 Educational Value and Trade-Offs", level=2)
    para(
        doc,
        "The hybrid design supports the original Flutter app requirement while making symbolic derivation practical. "
        "The trade-off is that a backend service is required for exact symbolic steps, but this keeps the frontend simpler and allows the mathematical rules to be tested more rigorously.",
    )

    doc.add_heading("Chapter 9: Conclusion and Future Work", level=1)
    doc.add_heading("9.1 Conclusion", level=2)
    para(
        doc,
        "This project developed a Flutter-based Fourier transform learning app with symbolic derivation and numerical visualisation. "
        "The project remains aligned with the original requirement for a smart-device Flutter application, while the Python backend provides supporting symbolic computation where Flutter-only processing is not sufficient. "
        "The result is a learning system that combines input interaction, examples, LaTeX display, teaching steps, and visual plots.",
    )

    doc.add_heading("9.2 Future Work", level=2)
    bullets(
        doc,
        [
            "Expand transform rule coverage and document active versus pending cases.",
            "Modularize the backend into rules, formatters, steps, API, and tests.",
            "Add screenshot or golden tests for visual regression.",
            "Improve user-facing unsupported-case messages.",
            "Add exportable derivation reports for classroom use.",
        ],
    )

    doc.add_heading("Appendix A: Reference Files", level=1)
    add_table(
        doc,
        ["Resource", "Path"],
        [
            ["Detailed Q&A", "docs/QA.md"],
            ["Brief email Q&A", "docs/QA_email_brief.md"],
            ["Dissertation outline", "docs/dissertation_detailed_outline.md"],
            ["Mapping from old report", "docs/word_to_dissertation_mapping.md"],
            ["Backend", "backend/backend.py"],
            ["Flutter home page", "flutter_app/lib/main.dart"],
            ["Symbolic result page", "flutter_app/lib/fft/symbol.dart"],
            ["FT Steps page", "flutter_app/lib/fft/step.dart"],
            ["Scrollable content widgets", "flutter_app/lib/scrollable_content.dart"],
        ],
        [3100, 6260],
    )

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
