from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Development_of_Flutter_Apps_for_Fourier_Transform_Dissertation.docx"


def _shade(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _cell_margins(cell, top=90, bottom=90, start=120, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("bottom", bottom), ("start", start), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def _table_width(table, widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), "9360")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[i]))
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            _cell_margins(cell)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        _shade(cell, "E8EEF5")
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value
    _table_width(table, widths)
    doc.add_paragraph()


def add_code(doc: Document, code: str) -> None:
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph(style="CodeBlock")
        run = p.add_run(line.rstrip())
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def para(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def setup(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.line_spacing = 1.15
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 14, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 10, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.bold = True
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("CodeBlock", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)


def title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(110)
    r = p.add_run("Development of Flutter Apps for Fourier Transform")
    r.bold = True
    r.font.size = Pt(25)
    r.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("With Symbolic Derivation and Numerical Visualisation: Supported by a Rule-First Python Backend")
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(46, 116, 181)

    for line in [
        "Dissertation Draft",
        "Flutter application as the main project focus",
        "Supporting backend used for symbolic Fourier derivation",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(line)
    doc.add_page_break()


def add_contents(doc: Document) -> None:
    doc.add_heading("Table of Contents", level=1)
    entries = [
        "1. Abstract",
        "2. Introduction",
        "3. Background and Design Rationale",
        "   - Fourier Transform Background",
        "   - Motivation",
        "   - Project Objectives",
        "   - Scope and Limitations",
        "4. Overall System Architecture",
        "   - System Overview",
        "   - Flutter Frontend",
        "   - FastAPI Symbolic Backend",
        "   - Frontend-Backend Communication",
        "5. Flutter Application Design",
        "   - User Interface",
        "   - Formula Input",
        "   - Example Navigation",
        "   - LaTeX Rendering",
        "6. Symbolic Fourier Transform Engine",
        "   - Definition-Based Framework",
        "   - Rule Matching Strategy",
        "   - Distribution-Based Transforms",
        "   - General Transform Families",
        "   - Error Handling",
        "7. Implementation Details",
        "   - Backend Commands and Functions",
        "   - Key Code Snippets",
        "   - Matcher Design",
        "   - API Design",
        "   - Data Flow",
        "8. Testing and Evaluation",
        "   - Backend Regression Tests",
        "   - Flutter Tests",
        "   - Representative Test Cases",
        "   - Performance Observations",
        "9. Discussion",
        "   - Current Capabilities",
        "   - Supported Function Classes",
        "   - Limitations",
        "   - Comparison with General CAS Systems",
        "10. Conclusion and Future Work",
        "References",
        "Appendix A. Question & Answer Summary",
        "Appendix B. Additional Code Snippets",
        "Appendix C. Example Inputs and Outputs",
    ]
    for entry in entries:
        doc.add_paragraph(entry)
    doc.add_page_break()


def build() -> None:
    doc = Document()
    setup(doc)
    title_page(doc)
    add_contents(doc)

    doc.add_heading("Abstract", level=1)
    para(doc, "This dissertation presents the development of a Flutter application for Fourier transform learning on smart devices. The original project requirement focuses on Flutter programming, simulator or device testing, step-by-step workings, and graphical illustrations of relevant Fourier transform equations. The project therefore keeps the Flutter application as the main user-facing system.")
    para(doc, "During development, a Flutter-only symbolic processing direction was considered. Flutter and Dart were suitable for user interaction, local numerical experiments, and visual presentation, but they were not sufficient for reliable symbolic Fourier transform derivation involving distribution-theory objects such as Dirac delta, Heaviside step functions, principal value terms, and engineering-style LaTeX steps. For this reason, the final design introduces a Python FastAPI backend as a supporting symbolic computation service.")
    para(doc, "The completed system combines a Flutter learning interface, structured examples, LaTeX display, numerical visualisation, and a rule-first backend for symbolic derivation. Current validation includes 68 passing backend tests and 6 passing Flutter tests, with 1 optional frontend-backend e2e test skipped unless explicitly enabled.")

    doc.add_heading("Introduction", level=1)
    para(doc, "The project was originally defined as the development of Flutter apps running on smart devices for Fourier transform. The expected application should be tested on a simulator or physical device, show step-by-step workings, and provide graph plots or illustrative graphics where useful.")
    para(doc, "This requirement makes the Flutter app the centre of the project. The backend introduced later should therefore be understood as a supporting service for symbolic derivation rather than as a replacement for the Flutter application objective.")
    para(doc, "The educational motivation is that students learning Fourier transform need to see both the final frequency-domain result and the reasoning that leads to it. A useful learning app should therefore support expression input, readable formulas, guided examples, derivation steps, and visual intuition.")

    doc.add_heading("Background and Design Rationale", level=1)
    doc.add_heading("Fourier Transform Background", level=2)
    para(doc, "The project uses the engineering Fourier transform convention")
    add_code(doc, "X(omega) = integral x(t) exp(-j omega t) dt")
    para(doc, "The display layer uses the engineering imaginary unit j and treats omega as real angular frequency. Distribution-theory objects such as delta, unit step, Cauchy principal value, and sign functions are used where ordinary integrals are not sufficient.")
    doc.add_heading("Motivation", level=2)
    para(doc, "General-purpose computer algebra systems can compute many symbolic transforms, but their output is often not aligned with engineering teaching notation. Results may contain Piecewise branches, RootSum expressions, complex assumptions, or other symbolic structures that are mathematically valid but difficult for students to interpret.")
    add_code(doc, """CAS-style issue:
Piecewise(... complex conditions ...)

Engineering textbook target:
F{1/(t^2+1)} = pi * exp(-abs(omega))""")
    doc.add_heading("Project Objectives", level=2)
    bullets(doc, [
        "Develop a Flutter frontend for expression input, examples, LaTeX display, and graph-based learning.",
        "Provide step-by-step workings for supported Fourier transform questions.",
        "Explain why a Flutter-only symbolic processing approach was not sufficient for the final derivation target.",
        "Use a Python backend to support distribution-aware symbolic derivation.",
        "Validate backend mathematical output and frontend layout behaviour with automated tests.",
    ])
    doc.add_heading("Scope and Limitations", level=2)
    para(doc, "The app is not designed as an unrestricted computer algebra system. It supports a defined engineering range, including constants, impulses, steps, trigonometric functions, exponentials, PV reciprocal forms, polynomial distributions, polynomial-step forms, selected rational forms, finite windows, and linear combinations of supported terms.")

    doc.add_heading("Overall System Architecture", level=1)
    doc.add_heading("System Overview", level=2)
    para(doc, "The system is best explained through responsibilities rather than a complex flow diagram. Flutter handles user-facing learning tasks, while the Python backend handles symbolic derivation. The two layers communicate through a JSON API.")
    add_table(doc, ["Layer", "Main responsibility"], [
        ["Flutter application", "Smart-device interface, keypad input, structured examples, LaTeX display, local numerical visualisation, and small-screen layout handling."],
        ["FastAPI symbolic backend", "Expression parsing, rule-first symbolic derivation, distribution theory rules, result_latex formatting, and steps_latex generation."],
        ["Communication boundary", "The Flutter app sends user expressions to /fourier and receives input_latex, result_latex, and steps_latex."],
    ], [2600, 6760])
    doc.add_heading("Flutter Frontend", level=2)
    para(doc, "The frontend is responsible for input, examples, formula preview, result display, numerical visualisation, and responsive layout. It remains the main application layer because the project requirement is specifically focused on Flutter apps for Fourier transform.")
    doc.add_heading("FastAPI Symbolic Backend", level=2)
    para(doc, "The backend is a supporting symbolic service. It was introduced because a reliable symbolic Fourier transform engine requires expression parsing, algebraic rewriting, distribution-theory rules, and controlled LaTeX output.")
    doc.add_heading("Frontend-Backend Communication", level=2)
    para(doc, "The frontend sends the actual user expression to the backend. This shows that the app is not limited to fixed examples; it can process general user input within the supported rule range.")
    add_code(doc, """final res = await http.post(
  uri,
  headers: {'Content-Type': 'application/json'},
  body: jsonEncode({'expression': expression}),
);""")

    doc.add_heading("Flutter Application Design", level=1)
    doc.add_heading("User Interface", level=2)
    para(doc, "The home page provides expression input, formula preview, example navigation, FFT controls, and access to symbolic and numerical result views.")
    doc.add_heading("Formula Input", level=2)
    para(doc, "The keypad and expression handling support common signal expressions such as u(t), delta(t), sin, cos, exp, powers, and fractions. The app also includes local numerical handling for visualisation.")
    doc.add_heading("Example Navigation", level=2)
    para(doc, "The example list was expanded from simple transform pairs into a structured set of teaching examples. Each example stores the expression, title, category, input LaTeX, transform object, and description.")
    add_code(doc, """class FourierExample {
  final String expression;
  final String title;
  final String category;
  final String inputLatex;
  final String transformLatex;
  final String description;
}""")
    doc.add_heading("LaTeX Rendering", level=2)
    para(doc, "The frontend uses LaTeX rendering for formula preview, symbolic results, derivation steps, and formula references. Long formulas and notes are handled through local scroll containers so that narrow screens do not cause overflow.")
    add_code(doc, """class ScrollableMathLine extends StatefulWidget {
  final String latex;
  final TextStyle? textStyle;
  final String? semanticsLabel;
}""")

    doc.add_heading("Symbolic Fourier Transform Engine", level=1)
    doc.add_heading("Definition-Based Framework", level=2)
    para(doc, "Every teaching derivation is grounded in the Fourier transform definition. The backend may use known pairs and properties, but the displayed derivation is written as a mathematical explanation rather than a trace of internal matcher logic.")
    doc.add_heading("Rule Matching Strategy", level=2)
    para(doc, "The backend applies rules before falling back to less controlled symbolic forms. This keeps the output predictable and closer to engineering textbook notation.")
    add_code(doc, """rules = [
    exact_matches,
    known_transform_pairs,
    distribution_rules,
    property_based_rules,
    decomposition_rules,
    controlled_fallback,
]""")
    doc.add_heading("Distribution-Based Transforms", level=2)
    para(doc, "Distribution-based transforms are a key reason for using a Python backend. For example, polynomial-step signals require delta derivatives and PV terms.")
    add_code(doc, """def _rule_poly_times_step_distribution(f):
    if not (isinstance(f, Mul) and f.has(Heaviside)):
        return None

    def _basis(n):
        if n == 0:
            return pi*DiracDelta(omega) - I*(1/omega)
        return (
            pi*(I**n)*Derivative(DiracDelta(omega), (omega, n))
            + ((-1)**n)*(I**(n-1))*factorial(n)*PV(1/(omega**(n+1)))
        )""")
    doc.add_heading("General Transform Families", level=2)
    para(doc, "The engine supports parameterized transform families rather than only fixed examples. These include shifted impulses, shifted steps, sinusoids with frequency and phase, complex exponentials, one-sided exponentials, PV reciprocal forms, polynomial distributions, polynomial-step forms, selected rational forms, finite windows, and linear combinations.")
    doc.add_heading("Error Handling", level=2)
    para(doc, "If an expression falls outside the supported range, the backend should return a controlled fallback or error instead of pretending to have a reliable closed-form result.")

    doc.add_heading("Implementation Details", level=1)
    doc.add_heading("Backend Commands and Functions", level=2)
    para(doc, "The backend symbolic API is implemented through the /fourier endpoint. The response contract explicitly separates the original input, final transform result, and derivation steps.")
    add_code(doc, """class FourierResponse(BaseModel):
    ok: bool
    input_latex: str
    result_latex: str
    steps_latex: list[str]
    error: str | None = None""")
    doc.add_heading("Key Code Snippets", level=2)
    para(doc, "The following function prepares derivation steps for teaching display by filtering internal implementation details and unwanted symbolic artifacts.")
    add_code(doc, """def _teaching_steps(f, X, steps):
    cleaned = []
    for raw in steps or []:
        s = _format_step_display_latex(str(raw).strip())
        if 'Method:' in s or '_rule_' in s or 'matcher' in s:
            continue
        if 'Piecewise' in s or 'RootSum' in s or 'meijerg' in s:
            continue
        cleaned.append(s)
    return cleaned""")
    doc.add_heading("Matcher Design", level=2)
    para(doc, "The matcher design is rule-first. This means exact and known transform families are attempted before fallback strategies. The design makes the supported range explicit and testable.")
    doc.add_heading("API Design", level=2)
    para(doc, "The API is intentionally simple: the frontend sends an expression string and receives LaTeX-ready fields. This design directly supports the Flutter UI and avoids exposing backend-specific symbolic objects to the frontend.")
    doc.add_heading("Data Flow", level=2)
    para(doc, "The data flow is: user expression in Flutter, POST request to /fourier, backend parsing and rule matching, LaTeX result generation, and Flutter rendering. This sequence is described in prose rather than as a diagram to keep the dissertation logic clear.")

    doc.add_heading("Testing and Evaluation", level=1)
    doc.add_heading("Backend Regression Tests", level=2)
    add_code(doc, """cd backend
.\\.venv\\Scripts\\python.exe -m pytest tests -q

Result: 68 passed""")
    para(doc, "Backend tests cover strict result_latex checks, teaching-style steps_latex checks, representative step snapshots, and output guards against Piecewise, RootSum, arg, polar_lift, meijerg, matcher, debug, and srepr.")
    doc.add_heading("Flutter Tests", level=2)
    add_code(doc, """cd flutter_app
flutter test

Result: 6 passed, 1 skipped""")
    para(doc, "Flutter tests cover responsive breakpoints, home page rendering, example navigation, multiple screen sizes without layout exceptions, local scroll containers for long formulas and notes, and FT Steps long-formula display on narrow screens.")
    doc.add_heading("Representative Test Cases", level=2)
    para(doc, "Representative backend step cases include 1, delta(t-3), u(t), t*u(t), sin(t)*u(t), frac(1,3*t-2), frac(1,t^2+1), and u(t-2)-u(t-5). These cases check both mathematical result quality and teaching-step clarity.")
    doc.add_heading("Performance Observations", level=2)
    para(doc, "Local numerical visualisation remains in Flutter for responsiveness, while symbolic derivation is handled by the backend. This separation avoids network latency for interactive chart updates and keeps exact symbolic work in the more suitable Python environment.")

    doc.add_heading("Discussion", level=1)
    doc.add_heading("Current Capabilities", level=2)
    para(doc, "The app supports general user input within a defined engineering range. It is not limited to fixed examples, although the example list helps guide users toward supported expression families.")
    doc.add_heading("Supported Function Classes", level=2)
    add_table(doc, ["Function class", "Examples"], [
        ["Impulses and steps", "delta(t), delta(t-a), u(t), u(t-a)"],
        ["Trigonometric and exponential", "sin(a*t+b), cos(a*t+b), exp(j*w0*t), exp(-a*t)u(t)"],
        ["Distribution and polynomial", "PV(1/t), 1/(a*t+b), t^n, t^n u(t)"],
        ["Rational and window forms", "1/(t^2+a^2), t/(t^2+a^2), u(t-a)-u(t-b)"],
    ], [3000, 6360])
    doc.add_heading("Limitations", level=2)
    para(doc, "The system is not an unrestricted CAS. Unsupported expressions may return an integral form or error. This limitation is acceptable because the project prioritises explainable and testable engineering output.")
    doc.add_heading("Comparison with General CAS Systems", level=2)
    para(doc, "General CAS systems prioritise symbolic completeness. This project prioritises educational readability, engineering notation, and step-by-step reasoning. The rule-first method provides cleaner output for the supported range but requires deliberate rule expansion.")

    doc.add_heading("Conclusion and Future Work", level=1)
    para(doc, "This project developed a Flutter-based Fourier transform learning app with symbolic derivation and numerical visualisation. The project remains aligned with the original Flutter smart-device app requirement, while the Python backend provides supporting symbolic computation where Flutter-only processing is not sufficient.")
    para(doc, "Future work includes expanding rule coverage, modularising the backend, improving unsupported-case explanations, adding screenshot-based visual regression tests, and supporting exportable derivation reports for classroom use.")

    doc.add_heading("References", level=1)
    bullets(doc, [
        "Course materials and standard Signals and Systems Fourier transform tables.",
        "Flutter documentation for application development and widget testing.",
        "FastAPI documentation for backend API implementation.",
        "SymPy documentation for symbolic parsing and algebraic manipulation.",
    ])

    doc.add_heading("Appendix A. Question & Answer Summary", level=1)
    para(doc, "A detailed Q&A file is maintained in docs/QA.md, and a shorter email version is maintained in docs/QA_email_brief.md. These files answer the supervisor's concerns about general user input capability, generic transform support within scope, PV explanation, examples, LaTeX display, derivation steps, testing, and dissertation formatting.")
    doc.add_heading("Appendix B. Additional Code Snippets", level=1)
    para(doc, "Longer implementation details, such as full rule functions, local numerical FFT code, parser details, and deployment configuration, should remain in the source repository or appendix rather than the main dissertation body.")
    doc.add_heading("Appendix C. Example Inputs and Outputs", level=1)
    para(doc, "The test data reference is maintained in test.md. It includes implemented transform cases and future/pending transform pairs for later rule expansion.")

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

