from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_ZH = ROOT / "傅里叶变换引擎_项目报告_精简修订版_v2.docx"
OUT_EN = ROOT / "Fourier_Transform_Engine_Project_Report_Concise_EN_v2.docx"


ACCENT = RGBColor(31, 78, 121)
MUTED = RGBColor(89, 89, 89)
LIGHT = RGBColor(242, 246, 250)


def set_cell_shading(cell, fill: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(9.5)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_width(table, widths: list[float]) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    tbl_pr = table._tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    total_dxa = int(sum(widths) * 1440)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(total_dxa))

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    for idx, width in enumerate(widths):
        table.columns[idx].width = Inches(width)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    p.style = f"Heading {level}"
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.color.rgb = ACCENT


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.style = "Body Text"
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.15
    run = p.add_run(text)
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run.font.size = Pt(10.5)


def add_kv_table(doc: Document, rows: list[tuple[str, str]], widths=(1.8, 4.9)) -> None:
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    set_table_width(table, list(widths))
    hdr = table.rows[0].cells
    set_cell_text(hdr[0], "项目项", True)
    set_cell_text(hdr[1], "当前说明", True)
    set_cell_shading(hdr[0], "D9EAF7")
    set_cell_shading(hdr[1], "D9EAF7")
    for key, value in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], key, True)
        set_cell_text(cells[1], value)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    table.autofit = False
    if widths:
        set_table_width(table, widths)
    for idx, header in enumerate(headers):
        cell = table.rows[0].cells[idx]
        set_cell_text(cell, header, True)
        set_cell_shading(cell, "D9EAF7")
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], text)
    doc.add_paragraph()


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.25)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def style_document(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"].font.size = Pt(10.5)
    for name, size in [("Heading 1", 16), ("Heading 2", 13), ("Heading 3", 11)]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = ACCENT
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)
    body = styles["Body Text"]
    body.font.name = "Microsoft YaHei"
    body.font.size = Pt(10.5)
    body.paragraph_format.line_spacing = 1.15


def build_zh() -> None:
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Fourier Transform Engine\n技术项目报告（精简修订版）")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("基于当前源码核对修订 | v2_fixed_20260121 | 2026-05")
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED

    add_heading(doc, "摘要", 1)
    add_body(
        doc,
        "本项目是一个面向信号与系统教学的傅里叶变换学习工具。系统采用 FastAPI 后端和 Flutter 前端：后端使用 SymPy 解析表达式，并通过规则优先的方式输出工程教材风格的傅里叶变换结果；前端负责公式输入、LaTeX 展示、本地数值近似和教学图表。与通用 CAS 相比，本项目更重视输出可读性、推导步骤和课堂演示体验。"
    )
    add_bullets(
        doc,
        [
            "核心目标：避免 RootSum、arg(ω) 和复杂 Piecewise，优先输出 δ、u(t)、sign(ω)、PV 等工程表达。",
            "当前形态：原型阶段，后端主要集中在 backend.py，前端主要集中在 Flutter 的 main.dart 与 fft/ 目录。",
            "修订重点：本版删去重复章节，修正 API 字段、前端 Tab 命名、规则优先级和部署说明，使报告与源码一致。"
        ],
    )

    add_heading(doc, "1. 项目定位", 1)
    add_body(
        doc,
        "通用符号计算系统追求数学完备性，但在教学场景中常产生冗长表达。本项目选择规则引擎而非完全依赖符号积分，目的是让结果更接近工程教材中的傅里叶变换表，并给出可直接渲染的分步 LaTeX 推导。"
    )
    add_kv_table(
        doc,
        [
            ("项目名称", "Fourier Transform Engine"),
            ("应用场景", "信号与系统课程演示、学生自学、习题验证、常见变换对查询"),
            ("后端技术", "FastAPI、Pydantic、SymPy、Starlette middleware"),
            ("前端技术", "Flutter、flutter_bloc、flutter_math_fork、Syncfusion Flutter Charts"),
            ("工程约定", "X(ω)=∫ x(t)e^{-jωt}dt，默认 ω 为实变量"),
        ],
        widths=(1.25, 4.85),
    )

    add_heading(doc, "2. 系统架构", 1)
    add_body(
        doc,
        "系统为前后端分离结构。后端只负责符号推导和工程化表达；前端负责输入、状态管理、数值采样、FFT 近似和图表展示。"
    )
    add_table(
        doc,
        ["层次", "主要文件", "职责"],
        [
            ["后端 API", "backend/backend.py", "定义 /fourier 接口、请求/响应模型、CORS 和性能追踪。"],
            ["符号规则", "backend/backend.py", "解析表达式，按规则优先级推导傅里叶变换，必要时回退到积分形式。"],
            ["前端入口", "flutter_app/lib/main.dart", "公式输入、示例切换、键盘和运行入口。"],
            ["状态与数值引擎", "flutter_app/lib/fft/fourier_transform_bloc.dart", "表达式本地解析、采样、复数运算、FFT 与卷积近似。"],
            ["结果页面", "flutter_app/lib/fft/charts.dart、symbol.dart、step.dart", "符号结果页和数值教学图表页。"],
        ],
        [1.15, 2.1, 3.0],
    )

    add_heading(doc, "3. 后端符号引擎", 1)
    add_body(
        doc,
        "后端暴露 POST /fourier。请求体当前只包含 expression 字段；full_apart 不是当前 API 模型的一部分。响应字段使用源码中的 Pydantic 模型：build_id、ok、input_latex、result_latex、steps_latex、error、method、form、conditions_latex。"
    )
    add_code_block(
        doc,
        '请求示例:\n{ "expression": "sin(t)*u(t)" }\n\n响应字段:\n{\n  "build_id": "...",\n  "ok": true,\n  "input_latex": "...",\n  "result_latex": "...",\n  "steps_latex": ["..."],\n  "error": null,\n  "method": "unknown",\n  "form": "distribution_form",\n  "conditions_latex": ""\n}'
    )
    add_body(
        doc,
        "注意：源码中虽然有 method 推断逻辑，但最终返回时 method 当前仍固定为 unknown。因此报告不应声称接口会稳定返回 trig_exp_distribution、exp_times_step 等具体规则名。若需要该能力，应先修正后端返回值。"
    )

    add_heading(doc, "4. 规则推导流程", 1)
    add_body(
        doc,
        "核心函数是 _derive_with_properties(f)。它不是一个完全模块化的规则注册系统，而是在单个函数中按顺序尝试规则和分支。当前优先处理 PV、纯多项式、t^n·u(t)、三角×阶跃、三角函数转复指数、有理函数、常数、δ、u(t)、指数×阶跃、纯音、线性组合和定义积分回退。"
    )
    add_table(
        doc,
        ["类型", "处理思路", "示例"],
        [
            ["常数/多项式", "按分布理论输出 δ 及其导数", "1、t、t^2"],
            ["δ 与 u(t)", "直接使用基本变换对和移位性质", "delta(t-a)、u(t-a)"],
            ["三角函数", "优先用欧拉公式转为复指数，再输出 δ 组合", "sin(t)、cos(3t+2)"],
            ["三角×阶跃", "使用 u(t) 的分布变换并做频移", "sin(t)*u(t)"],
            ["有理函数", "先 together/apart，再用实根部分分式作为备用方案", "1/(t^2+1)、(2t)/(3t^2+4t-1)"],
            ["卷积", "通过顶层卷积符号拆分左右表达式，并套用卷积定理", "sin(t)·cos(t)"],
        ],
        [1.05, 2.9, 1.7],
    )
    add_body(
        doc,
        "卷积符号需要特别说明：后端当前注释和解析逻辑以 ·（U+00B7）为主，而前端按钮使用 •。这属于实现与文档都需要统一的细节，建议统一为一个字符，并在 API 语法中明确。"
    )

    add_heading(doc, "5. 前端实现", 1)
    add_body(
        doc,
        "Flutter 前端使用 BLoC 管理状态。用户在首页选择示例或使用自定义键盘输入表达式，点击运行后，本地数值路径会先完成采样和 FFT 近似；结果页再通过 SymbolPage 请求后端获得符号推导。"
    )
    add_table(
        doc,
        ["页面/组件", "当前行为"],
        [
            ["HomePage", "显示公式预览、示例切换、符号键盘、FFT size 滑块和运行按钮。"],
            ["ResultsPage", "当前只有两个 Tab：Step 与 Chart；第一个 Tab 实际承载 SymbolPage，第二个 Tab 承载 StepPage。"],
            ["SymbolPage", "通过 HTTP 调用后端 /fourier，展示输入、结果和推导步骤。"],
            ["StepPage", "展示原信号、核函数、被积函数、累计积分和 ω 扫描曲线。"],
            ["本地数值引擎", "支持复数计算、FFT、圆周卷积、frac、sin/cos/exp/u/delta 等基础表达式。"],
        ],
        [1.35, 4.65],
    )
    add_body(
        doc,
        "当前 UI 有一个小不一致：代码保留了 t^ 插入函数，但 abs 按钮实际绑定的是 abs 插入逻辑。若报告要描述专业符号键盘，应以当前界面为准，或先补回 t^ 按钮。"
    )

    add_heading(doc, "6. 教学可视化", 1)
    add_body(
        doc,
        "数值可视化不试图替代符号结果，而是解释傅里叶积分的过程。StepPage 以 Riemann 和近似展示 x(t)、e^{-jωt}、x(t)e^{-jωt}、累计积分和 ω 轴扫描结果。该设计适合课堂中说明“频率匹配时相关性增强”的直观含义。"
    )
    add_bullets(
        doc,
        [
            "原信号图：展示采样窗口内的 x(t)。",
            "核函数图：展示 cos(ωt) 与 -sin(ωt)。",
            "被积函数图：展示实部和虚部积分项。",
            "累计积分图：展示从左到右累加时 X(ω) 的收敛过程。",
            "ω 扫描图：对 201 个 ω 值计算近似积分，展示频谱趋势。"
        ],
    )

    add_heading(doc, "7. 性能与工程实践", 1)
    add_table(
        doc,
        ["主题", "当前实现", "建议"],
        [
            ["后端性能", "对 together 和 apart 使用 srepr 字典缓存。", "后续可改为 LRU 缓存并限制容量。"],
            ["慢请求追踪", "使用 contextvars 和中间件记录慢操作，默认关闭。", "保留默认关闭，用环境变量控制更合适。"],
            ["代码组织", "backend.py 仍是单文件集中实现。", "拆分为 api、parser、rules、postprocess、cache。"],
            ["测试", "当前报告声称生产可用，但仓库未见系统化测试。", "补 pytest 和 Flutter widget/unit tests 后再使用“生产可用”措辞。"],
            ["部署配置", "CORS 只允许 GitHub Pages，前端 Web 默认请求 Render。", "用环境变量配置 API_BASE_URL 和 ALLOWED_ORIGINS。"],
        ],
        [1.1, 2.25, 2.75],
    )

    add_heading(doc, "8. 当前限制", 1)
    add_bullets(
        doc,
        [
            "覆盖范围以课程常见信号为主，不等同于通用 CAS。",
            "method 字段尚未返回真实规则名。",
            "卷积符号在前后端和文档中需要统一。",
            "符号结果与数值图表是两条路径，数值路径对 δ、PV、非绝对可积信号只能做近似展示。",
            "中文/特殊符号在部分源码注释和旧 README 中曾出现编码损坏，正式发布前应统一 UTF-8。"
        ],
    )

    add_heading(doc, "9. 后续计划", 1)
    add_table(
        doc,
        ["优先级", "事项", "价值"],
        [
            ["高", "修正 API 文档与 method 返回值", "降低前后端联调和第三方接入成本。"],
            ["高", "统一卷积符号", "避免用户输入与后端解析不一致。"],
            ["中", "拆分 backend.py", "提升规则维护性和测试便利性。"],
            ["中", "补充自动化测试", "验证常见变换对和回归风险。"],
            ["低", "扩展 sinc、Bessel、逆变换等规则", "扩大课程和工程覆盖范围。"],
        ],
        [0.85, 2.2, 3.0],
    )

    add_heading(doc, "结论", 1)
    add_body(
        doc,
        "Fourier Transform Engine 已经具备清晰的教学价值：它把工程傅里叶变换表、分布理论和可视化积分过程结合到一个可交互工具中。当前报告应避免过度宣称覆盖率和生产状态，更准确地定位为“规则驱动的教学原型”。在修正 API 契约、统一符号输入、补充测试和拆分模块后，项目可以继续向稳定教学工具演进。"
    )

    doc.add_section(WD_SECTION.CONTINUOUS)
    footer = doc.sections[-1].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = footer.add_run("Fourier Transform Engine 技术项目报告（精简修订版）")
    rr.font.name = "Microsoft YaHei"
    rr.font.size = Pt(8)
    rr.font.color.rgb = MUTED

    doc.save(OUT_ZH)
    print(OUT_ZH)


def build_en() -> None:
    doc = Document()
    style_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Fourier Transform Engine\nTechnical Project Report (Concise Edition)")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run.font.size = Pt(20)
    run.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Revised against the current source code | v2_fixed_20260121 | May 2026")
    r.font.name = "Microsoft YaHei"
    r.font.size = Pt(10)
    r.font.color.rgb = MUTED

    add_heading(doc, "Executive Summary", 1)
    add_body(
        doc,
        "Fourier Transform Engine is a teaching-oriented Fourier transform tool for signals and systems courses. It combines a FastAPI backend with a Flutter frontend. The backend parses expressions with SymPy and applies rule-first symbolic derivation to produce engineering-style results. The frontend provides formula input, LaTeX rendering, local numerical approximation, and interactive educational charts."
    )
    add_bullets(
        doc,
        [
            "Goal: prefer textbook-style expressions using δ, u(t), sign(ω), and PV, while avoiding RootSum, arg(ω), and overly complex Piecewise output.",
            "Current state: a functional teaching prototype, with the backend concentrated in backend.py and the frontend concentrated in main.dart plus the fft/ directory.",
            "Revision focus: this edition removes duplicate generated sections and aligns the API, frontend page structure, rule flow, and deployment notes with the current repository."
        ],
    )

    add_heading(doc, "1. Project Positioning", 1)
    add_body(
        doc,
        "General-purpose computer algebra systems aim for mathematical completeness, but their Fourier transform output can be difficult to use in a classroom. This project deliberately favors a rule-based engine over raw symbolic integration so that the output resembles engineering transform tables and includes readable LaTeX derivation steps."
    )
    add_kv_table(
        doc,
        [
            ("Project", "Fourier Transform Engine"),
            ("Use cases", "Course demonstrations, self-study, exercise checking, and common transform lookup"),
            ("Backend stack", "FastAPI, Pydantic, SymPy, Starlette middleware"),
            ("Frontend stack", "Flutter, flutter_bloc, flutter_math_fork, Syncfusion Flutter Charts"),
            ("Convention", "X(ω)=∫ x(t)e^{-jωt}dt, with ω treated as real by default"),
        ],
        widths=(1.25, 4.85),
    )

    add_heading(doc, "2. System Architecture", 1)
    add_body(
        doc,
        "The system uses a separated frontend/backend architecture. The backend handles symbolic derivation and engineering-style formatting. The frontend handles input, state management, local sampling, FFT approximation, and chart rendering."
    )
    add_table(
        doc,
        ["Layer", "Main Files", "Responsibility"],
        [
            ["Backend API", "backend/backend.py", "Defines /fourier, request/response models, CORS, and performance tracing."],
            ["Symbolic rules", "backend/backend.py", "Parses expressions and applies rule-first Fourier derivation with fallback integral forms."],
            ["Frontend entry", "flutter_app/lib/main.dart", "Formula input, examples, keypad, and run action."],
            ["State and numeric engine", "flutter_app/lib/fft/fourier_transform_bloc.dart", "Local parsing, sampling, complex arithmetic, FFT, and convolution approximation."],
            ["Result pages", "charts.dart, symbol.dart, step.dart", "Symbolic result display and numerical teaching charts."],
        ],
        [1.2, 2.25, 2.75],
    )

    add_heading(doc, "3. Backend Symbolic Engine", 1)
    add_body(
        doc,
        "The backend exposes POST /fourier. The current request body contains only the expression field; full_apart is not part of the current API model. The response model contains build_id, ok, input_latex, result_latex, steps_latex, error, method, form, and conditions_latex."
    )
    add_code_block(
        doc,
        'Request:\n{ "expression": "sin(t)*u(t)" }\n\nResponse fields:\n{\n  "build_id": "...",\n  "ok": true,\n  "input_latex": "...",\n  "result_latex": "...",\n  "steps_latex": ["..."],\n  "error": null,\n  "method": "unknown",\n  "form": "distribution_form",\n  "conditions_latex": ""\n}'
    )
    add_body(
        doc,
        "Important implementation note: although the backend contains method inference logic, the final response currently returns method as unknown. The report should not claim stable rule names such as trig_exp_distribution or exp_times_step until the backend returns them explicitly."
    )

    add_heading(doc, "4. Rule Derivation Flow", 1)
    add_body(
        doc,
        "The central function is _derive_with_properties(f). It is not yet a fully modular rule registry. Instead, a single function attempts rules and branches in a fixed order. The current flow prioritizes PV rules, polynomials, t^n·u(t), trigonometric step signals, trigonometric-to-exponential expansion, rational functions, constants, δ, u(t), exponential step signals, pure tones, linearity, and fallback integrals."
    )
    add_table(
        doc,
        ["Type", "Approach", "Examples"],
        [
            ["Constants and polynomials", "Use distribution theory and derivatives of δ.", "1, t, t^2"],
            ["δ and u(t)", "Apply known pairs and time-shift properties.", "delta(t-a), u(t-a)"],
            ["Trigonometric signals", "Rewrite with Euler identities and return δ combinations.", "sin(t), cos(3t+2)"],
            ["Trig × step", "Combine the u(t) distribution pair with frequency shifts.", "sin(t)*u(t)"],
            ["Rational functions", "Use together/apart first, then real-root partial fractions as fallback.", "1/(t^2+1)"],
            ["Convolution", "Split on the top-level convolution symbol and apply the convolution theorem.", "sin(t)·cos(t)"],
        ],
        [1.25, 3.2, 1.45],
    )
    add_body(
        doc,
        "The convolution symbol should be standardized. The backend currently emphasizes · (U+00B7), while the frontend keypad uses •. One symbol should be chosen and reflected consistently in code, UI, and API documentation."
    )

    add_heading(doc, "5. Frontend Implementation", 1)
    add_body(
        doc,
        "The Flutter application uses BLoC for state management. After the user selects an example or enters an expression and runs it, the local numeric path computes samples and an FFT approximation. The result screen then uses SymbolPage to call the backend for symbolic derivation."
    )
    add_table(
        doc,
        ["Component", "Current Behavior"],
        [
            ["HomePage", "Formula preview, example navigation, custom keypad, FFT size slider, and run button."],
            ["ResultsPage", "Two tabs: Step and Chart. The first tab currently hosts SymbolPage; the second hosts StepPage."],
            ["SymbolPage", "Calls /fourier over HTTP and displays input, result, and derivation steps."],
            ["StepPage", "Shows original signal, kernel, integrand, cumulative integral, and ω sweep chart."],
            ["Local numeric engine", "Supports complex arithmetic, FFT, circular convolution, frac, sin/cos/exp/u/delta, and related basics."],
        ],
        [1.55, 4.55],
    )
    add_body(
        doc,
        "There is a minor UI mismatch: the code still contains a t^ insertion helper, but the visible abs button is currently wired to insert abs(...). The report should describe the current UI, or the UI should restore a dedicated t^ button."
    )

    add_heading(doc, "6. Educational Visualization", 1)
    add_body(
        doc,
        "The numerical visualization complements the symbolic result. It does not replace the exact backend derivation. StepPage uses Riemann sums to show x(t), e^{-jωt}, x(t)e^{-jωt}, the cumulative integral, and an ω-axis sweep. This is useful for explaining how frequency matching leads to stronger correlation."
    )
    add_bullets(
        doc,
        [
            "Original signal: the sampled x(t) over the selected window.",
            "Kernel: cos(ωt) and -sin(ωt).",
            "Integrand: real and imaginary integrand components.",
            "Cumulative integral: the running approximation of X(ω).",
            "ω sweep: an approximate spectrum computed over 201 frequency points."
        ],
    )

    add_heading(doc, "7. Performance and Engineering Notes", 1)
    add_table(
        doc,
        ["Topic", "Current Implementation", "Recommendation"],
        [
            ["Backend performance", "srepr-based dictionaries cache together and apart results.", "Replace with bounded LRU caches."],
            ["Tracing", "contextvars-based slow-operation tracing is present and disabled by default.", "Control it through environment variables."],
            ["Code organization", "backend.py is still a single large file.", "Split into api, parser, rules, postprocess, and cache modules."],
            ["Testing", "The repository does not show systematic automated tests.", "Add pytest and Flutter unit/widget tests before claiming production readiness."],
            ["Deployment config", "CORS allows GitHub Pages, while Flutter Web defaults to a Render backend URL.", "Move API_BASE_URL and ALLOWED_ORIGINS to environment configuration."],
        ],
        [1.35, 2.45, 2.65],
    )

    add_heading(doc, "8. Current Limitations", 1)
    add_bullets(
        doc,
        [
            "The coverage targets common course signals and should not be presented as a general CAS replacement.",
            "The method field does not yet report actual rule names.",
            "The convolution symbol must be standardized across backend, frontend, and documentation.",
            "Symbolic results and numerical charts are separate paths; δ, PV, and non-absolutely-integrable signals can only be approximated visually.",
            "Some older comments and README text show signs of encoding corruption and should be normalized to UTF-8 before release."
        ],
    )

    add_heading(doc, "9. Roadmap", 1)
    add_table(
        doc,
        ["Priority", "Item", "Value"],
        [
            ["High", "Fix API documentation and method return behavior", "Improves integration reliability."],
            ["High", "Standardize the convolution symbol", "Prevents user input and backend parsing mismatch."],
            ["Medium", "Modularize backend.py", "Improves rule maintainability and testability."],
            ["Medium", "Add automated tests", "Reduces regression risk across transform pairs."],
            ["Low", "Add sinc, Bessel, inverse transform, and related rules", "Expands educational and engineering coverage."],
        ],
        [0.9, 2.65, 2.45],
    )

    add_heading(doc, "Conclusion", 1)
    add_body(
        doc,
        "Fourier Transform Engine already has clear educational value: it combines engineering transform tables, distribution theory, and interactive visualization in one tool. The most accurate current positioning is a rule-driven teaching prototype. After API cleanup, symbol unification, tests, and backend modularization, it can evolve into a more stable classroom-ready system."
    )

    doc.add_section(WD_SECTION.CONTINUOUS)
    footer = doc.sections[-1].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = footer.add_run("Fourier Transform Engine Technical Project Report (Concise Edition)")
    rr.font.name = "Microsoft YaHei"
    rr.font.size = Pt(8)
    rr.font.color.rgb = MUTED

    doc.save(OUT_EN)
    print(OUT_EN)


if __name__ == "__main__":
    build_zh()
    build_en()
