from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "Development_of_Flutter_Apps_for_Fourier_Transform_Dissertation_Draft.docx"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int, col_widths: list[int]) -> None:
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_ind = OxmlElement("w:tblInd")
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_pr.append(tbl_ind)

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.append(grid)
    for child in list(grid):
        grid.remove(child)
    for w in col_widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(w))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.tcW
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(col_widths[idx]))
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_code(doc: Document, code: str) -> None:
    for line in code.strip("\n").splitlines():
        p = doc.add_paragraph(style="CodeBlock")
        run = p.add_run(line.rstrip())
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")


def add_note(doc: Document, label: str, text: str) -> None:
    p = doc.add_paragraph(style="Callout")
    r = p.add_run(label + ": ")
    r.bold = True
    p.add_run(text)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        set_cell_shading(cell, "E8EEF5")
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value
    set_table_width(table, 9360, widths)
    doc.add_paragraph()


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color, before, after in [
        ("Heading 1", 16, RGBColor(46, 116, 181), 14, 8),
        ("Heading 2", 13, RGBColor(46, 116, 181), 10, 6),
        ("Heading 3", 12, RGBColor(31, 77, 120), 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    code = styles.add_style("CodeBlock", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(9)
    code.paragraph_format.space_before = Pt(0)
    code.paragraph_format.space_after = Pt(0)
    code.paragraph_format.left_indent = Inches(0.18)

    callout = styles.add_style("Callout", 1)
    callout.font.name = "Calibri"
    callout.font.size = Pt(10.5)
    callout.paragraph_format.left_indent = Inches(0.18)
    callout.paragraph_format.right_indent = Inches(0.18)
    callout.paragraph_format.space_before = Pt(6)
    callout.paragraph_format.space_after = Pt(6)


def title_page(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    r = p.add_run("Development of Flutter Apps for Fourier Transform")
    r.bold = True
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor(11, 37, 69)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(
        "With Symbolic Derivation and Numerical Visualisation: "
        "Supported by a Rule-First Python Backend"
    )
    r.font.size = Pt(15)
    r.font.color.rgb = RGBColor(46, 116, 181)

    for label in [
        "Dissertation Draft",
        "Prepared from the original Flutter application project requirement and recent implementation updates",
        "Main application: Flutter | Supporting symbolic service: FastAPI + SymPy",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run(label)

    doc.add_page_break()


def build() -> None:
    doc = Document()
    setup_styles(doc)
    title_page(doc)

    doc.add_heading("Abstract", level=1)
    doc.add_paragraph(
        "This dissertation presents the development of a Flutter application for Fourier transform learning on smart devices. "
        "The original project requirement focuses on Flutter programming, simulator or device testing, step-by-step workings, and graphical illustration of Fourier transform concepts. "
        "The project therefore keeps Flutter as the main user-facing application while extending it with a supporting Python backend for reliable symbolic derivation."
    )
    doc.add_paragraph(
        "During early development, local Flutter/Dart-side symbolic processing was considered for the entire application. "
        "This approach was suitable for expression input, numerical experimentation, and visual interaction, but it was not sufficient for robust symbolic Fourier transform derivation involving distribution-theory objects such as Dirac delta, Heaviside step functions, principal value terms, and controlled engineering-style LaTeX output. "
        "For this reason, the final system adopts a hybrid architecture: Flutter provides the learning interface and visualisation, while a FastAPI/SymPy backend provides rule-first symbolic computation."
    )
    doc.add_paragraph(
        "Recent refinements include clearer teaching-step generation, consistent engineering display conventions using j, u(t), PV, sign, compact absolute-value notation, and cleaner delta derivative formatting. "
        "The frontend was also improved to handle long formulas and annotations on small screens through local scroll containers. "
        "Validation currently includes backend regression tests with 68 passing cases and Flutter tests with 6 passing cases and 1 intentionally skipped optional e2e test. "
        "The system is not intended to be an unrestricted CAS; it supports a defined and expandable engineering range."
    )

    doc.add_heading("Chapter 1: Introduction", level=1)
    doc.add_heading("1.1 Background", level=2)
    doc.add_paragraph(
        "This project is based on the requirement to develop Flutter apps for Fourier transform on smart devices. "
        "The application is intended for learning and teaching: users should be able to enter Fourier transform questions, view step-by-step workings, and inspect relevant graph plots or illustrative graphics where possible. "
        "Fourier transform methods are central to signals and systems, communications, control engineering, and signal processing, so a mobile-friendly educational interface can support both classroom demonstration and self-study."
    )

    doc.add_heading("1.2 Problem Statement", level=2)
    doc.add_paragraph(
        "General-purpose CAS tools are optimized for symbolic completeness, not necessarily for pedagogical clarity. "
        "For example, a transform that is usually written in textbooks as a compact expression may be returned by a CAS using conditions, branch structures, or special symbolic wrappers. "
        "This makes the result difficult to compare with standard engineering transform tables."
    )
    add_code(
        doc,
        r"""
Engineering textbook form:
F{1/(t^2 + 1)} = pi * exp(-abs(omega))
""",
    )
    doc.add_paragraph(
        "The project therefore focuses on producing textbook-like Fourier transform outputs and derivation steps within a defined range of common engineering signals, while keeping Flutter as the main application layer."
    )

    doc.add_heading("1.3 Aim and Objectives", level=2)
    doc.add_paragraph(
        "The aim of this project is to design and implement a Flutter-based Fourier transform learning application that runs on smart-device-oriented platforms and provides symbolic derivation, numerical visualisation, and teaching-oriented explanations for a defined range of common signals."
    )
    for item in [
        "Build a Flutter frontend for expression input, examples, LaTeX display, and graph-based learning.",
        "Evaluate the practicality of local Flutter/Dart symbolic processing and define a suitable frontend/backend boundary.",
        "Build a rule-first Python backend symbolic transform engine where local Flutter symbolic processing is insufficient.",
        "Support distribution-aware results involving delta, unit step, PV, and sign terms.",
        "Generate readable result_latex and steps_latex outputs.",
        "Provide a Flutter frontend for expression input, examples, LaTeX preview, symbolic results, and numerical visualisation.",
        "Improve small-screen usability for long formulas and annotations.",
        "Validate the system with backend and frontend automated tests.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("1.4 Scope and Contributions", level=2)
    doc.add_paragraph(
        "The system is not designed as an unrestricted symbolic mathematics system. "
        "Instead, it supports a defined engineering range that can be expanded through additional rules. "
        "This scope is appropriate for an educational tool because predictable, explainable results are more valuable than uncontrolled symbolic complexity."
    )
    add_table(
        doc,
        ["Contribution", "Description"],
        [
            ["Flutter learning app", "Provides the main smart-device-oriented user interface for input, examples, formulas, and plots."],
            ["Rule-first backend", "Supports symbolic derivation where Flutter-only symbolic processing is not reliable enough."],
            ["Distribution-aware output", "Uses delta, unit step, PV, and sign terms in engineering notation."],
            ["Teaching steps", "Generates derivations that begin from the Fourier definition and explain selected transform pairs or properties."],
            ["Flutter learning interface", "Provides examples, LaTeX display, symbolic results, and numerical visualisation."],
            ["Responsive formula display", "Uses local scroll containers for long formulas and notes on small screens."],
            ["Automated validation", "Backend and frontend tests guard mathematical results, notation, steps, and layout behavior."],
        ],
        [2700, 6660],
    )

    doc.add_heading("Chapter 2: Mathematical Background", level=1)
    doc.add_heading("2.1 Fourier Transform Convention", level=2)
    add_code(
        doc,
        r"""
X(omega) = integral_{-infinity}^{infinity} x(t) * exp(-j*omega*t) dt
""",
    )
    doc.add_paragraph(
        "The project uses the engineering Fourier transform convention with angular frequency omega and imaginary unit j. "
        "The output layer prioritises frequency-domain expressions in distribution form where appropriate."
    )

    doc.add_heading("2.2 Distribution Theory Objects", level=2)
    add_table(
        doc,
        ["Object", "Role in the project"],
        [
            ["Dirac delta", "Represents impulses in the frequency domain, such as the transform of constants or sinusoids."],
            ["Unit step u(t)", "Represents one-sided signals and produces delta plus PV terms."],
            ["PV", "Represents Cauchy principal value terms for non-absolutely integrable signals."],
            ["sign(omega)", "Appears in principal value reciprocal transforms."],
        ],
        [2300, 7060],
    )
    add_code(
        doc,
        r"""
F{u(t)} = pi*delta(omega) - j*PV(1/omega)
F{PV(1/t)} = -j*pi*sign(omega)
""",
    )

    doc.add_heading("2.3 Transform Pairs and Properties", level=2)
    add_table(
        doc,
        ["Family", "Example"],
        [
            ["Impulse", "delta(t), delta(t-a)"],
            ["Step", "u(t), u(t-a)"],
            ["Trigonometric", "sin(a*t+b), cos(a*t+b)"],
            ["Exponential", "exp(j*w0*t), exp(-a*t)u(t), exp(-a*abs(t))"],
            ["Polynomial distribution", "t^n"],
            ["Polynomial-step distribution", "t^n u(t)"],
            ["PV reciprocal", "1/t, 1/(a*t+b)"],
            ["Rational", "1/(t^2+a^2), t/(t^2+a^2), selected quadratic forms"],
        ],
        [2800, 6560],
    )

    doc.add_heading("Chapter 3: System Requirements and Design Rationale", level=1)
    doc.add_heading("3.1 Requirements", level=2)
    for item in [
        "Accept user-entered mathematical expressions and preset examples.",
        "Return symbolic transform results using LaTeX.",
        "Return teaching-style derivation steps.",
        "Display formulas clearly on desktop and small screens.",
        "Keep frontend deployment independent from backend deployment.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("3.2 Design Rationale", level=2)
    doc.add_paragraph(
        "The central design decision is to use a rule-first transform engine rather than pure symbolic integration. "
        "The backend still uses SymPy for parsing and algebraic manipulation, but it avoids using integration as the main route because the result style cannot be reliably controlled. "
        "The rule-first design allows the engine to select a known transform pair, apply a property, and generate a matching teaching explanation."
    )
    add_note(
        doc,
        "Trade-off",
        "The rule-first approach improves explainability and output control, but coverage is limited to implemented rule families.",
    )

    doc.add_heading("3.3 From Flutter-Only Processing to Python Backend Support", level=2)
    doc.add_paragraph(
        "An early design option was to implement both the application interface and the symbolic Fourier transform logic locally in Flutter/Dart. "
        "This would have simplified deployment because the app could run without a separate backend service. "
        "The Flutter side was suitable for interface construction, keypad input, formula preview, numerical evaluation, and graph-based visualisation."
    )
    doc.add_paragraph(
        "However, local Flutter-only symbolic processing was found to be unsuitable for the level of derivation required by the project. "
        "Symbolic Fourier transform derivation requires robust expression parsing, algebraic rewriting, transform-pair matching, distribution-theory objects, and controlled LaTeX output. "
        "Implementing these features from scratch in Dart would require a large custom symbolic engine and would still be difficult to validate."
    )
    doc.add_paragraph(
        "The final design therefore adopts a hybrid architecture. "
        "Flutter remains the main learning application and handles user interaction, examples, visualisation, and responsive display. "
        "The Python backend acts as a supporting symbolic computation service, generating engineering-style Fourier transform results and teaching-oriented derivation steps."
    )

    doc.add_heading("3.4 Architecture", level=2)
    add_code(
        doc,
        r"""
Flutter UI
  -> expression input / examples / LaTeX preview
  -> POST /fourier
FastAPI backend
  -> parser
  -> rule-first transform engine
  -> LaTeX formatter
  -> teaching step generator
Flutter UI
  -> x(t), X(omega), and steps_latex display
  -> numerical FT visualisation
""",
    )
    doc.add_paragraph(
        "The frontend is designed to build independently, for example through GitHub Actions, while the backend can be deployed separately on Render. "
        "The API base URL can be configured at build time using a Dart define."
    )

    doc.add_heading("Chapter 4: Backend Symbolic Fourier Transform Engine", level=1)
    doc.add_heading("4.1 API Contract", level=2)
    add_code(
        doc,
        r"""
class FourierRequest(BaseModel):
    expression: str

class FourierResponse(BaseModel):
    ok: bool
    input_latex: str
    result_latex: str
    steps_latex: list[str]
    error: str | None = None
    method: str | None = None
    form: str | None = None
    conditions_latex: str | None = None
""",
    )
    doc.add_paragraph(
        "The response separates the final transform result from the derivation steps. "
        "This allows the frontend to display both a compact answer and a teaching explanation."
    )

    doc.add_heading("4.2 Rule-First Pipeline", level=2)
    for item in [
        "Parse and normalise the input expression.",
        "Try direct transform-pair rules.",
        "Try distribution-aware rules.",
        "Apply Fourier properties such as shift, modulation, scale, and linearity.",
        "Generate controlled LaTeX result and steps.",
        "Return an integral form or error if a closed-form rule is unavailable.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("4.3 Distribution Rule Example: Polynomial Times Step", level=2)
    add_code(
        doc,
        r"""
def _rule_poly_times_step_distribution(f):
    if not (isinstance(f, Mul) and f.has(Heaviside)):
        return None

    g = expand(simplify(f / h_t))
    if not g.is_polynomial(t):
        return None

    def _basis(n):
        if n == 0:
            return pi*DiracDelta(omega) - I*(1/omega)
        return (
            pi*(I**n)*Derivative(DiracDelta(omega), (omega, n))
            + ((-1)**n)*(I**(n-1))*factorial(n)*PV(1/(omega**(n+1)))
        )
""",
    )
    doc.add_paragraph(
        "This rule recognizes polynomial-step signals and applies the distribution formula term by term. "
        "It is a representative example of how the backend derives structured engineering results rather than delegating everything to symbolic integration."
    )

    doc.add_heading("4.4 Teaching Step Generation", level=2)
    add_code(
        doc,
        r"""
def _teaching_steps(f, X, steps):
    cleaned = []
    for raw in steps or []:
        s = _format_step_display_latex(str(raw).strip())
        if "Method:" in s or "_rule_" in s or "matcher" in s:
            continue
        if "Piecewise" in s or "RootSum" in s or "meijerg" in s:
            continue
        cleaned.append(s)

    if r"X(\omega)=\int" not in "\n".join(cleaned):
        cleaned = _step_start_definition(f) + cleaned

    if "Final Result" not in "\n".join(cleaned):
        cleaned += [r"\textbf{Final Result}", r"X(\omega)=" + result_text]
    return cleaned
""",
    )
    doc.add_paragraph(
        "The step generator removes internal implementation details and ensures that the displayed derivation begins with the Fourier transform definition. "
        "It also filters out unwanted symbolic artifacts and adds a final result if needed."
    )

    doc.add_heading("4.5 Engineering Output Formatting", level=2)
    doc.add_paragraph(
        "Recent refinements made the output notation more consistent with engineering practice. "
        "The display layer uses j rather than i, shows Heaviside/theta as u(t), formats principal value terms as PV fractions, simplifies sign and absolute value notation, and uses compact delta derivative notation."
    )

    doc.add_heading("Chapter 5: Frontend Design and Implementation", level=1)
    doc.add_heading("5.1 Home Page and Examples", level=2)
    doc.add_paragraph(
        "The Flutter home page provides keypad-based expression input, formula preview, FFT controls, and categorized examples. "
        "The example list was expanded to include shift, scale, combination, PV, rational, polynomial-step, and finite-window cases."
    )
    add_code(
        doc,
        r"""
class FourierExample {
  final String expression;
  final String title;
  final String category;
  final String inputLatex;
  final String transformLatex;
  final String description;
}
""",
    )

    doc.add_heading("5.2 Symbolic Result Page", level=2)
    add_code(
        doc,
        r"""
Future<SymbolicResult> computeByBackendOnly(String expression) async {
  const String envBase = String.fromEnvironment(
    'API_BASE_URL',
    defaultValue: '',
  );

  final Uri uri = Uri.parse(
    base.endsWith('/') ? '${base}fourier' : '${base}/fourier',
  );

  final res = await http.post(
    uri,
    headers: {'Content-Type': 'application/json'},
    body: jsonEncode({'expression': expression}),
  );
}
""",
    )
    doc.add_paragraph(
        "This code shows that the frontend can target different backend locations without changing the source code. "
        "It supports the current deployment model in which the frontend and backend are built separately."
    )

    doc.add_heading("5.3 Formula Rendering and Small-Screen Handling", level=2)
    doc.add_paragraph(
        "Long formulas and long annotations can overflow on narrow devices. "
        "The frontend now uses local scroll containers so that only the long formula or note area scrolls, while the surrounding page remains stable."
    )
    add_code(
        doc,
        r"""
class ScrollableMathLine extends StatefulWidget {
  final String latex;
  final TextStyle? textStyle;
  final String? semanticsLabel;
}

class ScrollableTextLine extends StatefulWidget {
  final String text;
  final TextStyle? style;
  final String? semanticsLabel;
}

class BoundedScrollableText extends StatefulWidget {
  final String text;
  final TextStyle? style;
  final double maxHeight;
}
""",
    )

    doc.add_heading("Chapter 6: Numerical Fourier Transform and Teaching Visualisation", level=1)
    doc.add_heading("6.1 Purpose", level=2)
    doc.add_paragraph(
        "The symbolic backend answers why a transform has a particular form, while the numerical frontend helps students see what the transform process does. "
        "The numerical page therefore complements the symbolic engine rather than replacing it."
    )
    doc.add_heading("6.2 FT Steps Page", level=2)
    for item in [
        "Original signal x(t).",
        "Kernel e^{-j omega t}.",
        "Integrand x(t)e^{-j omega t}.",
        "Running integral as cumulative summation.",
        "Frequency scan curve over omega.",
    ]:
        doc.add_paragraph(item, style="List Number")

    doc.add_heading("Chapter 7: Testing and Validation", level=1)
    doc.add_heading("7.1 Backend Tests", level=2)
    add_code(
        doc,
        r"""
cd backend
.\.venv\Scripts\python.exe -m pytest tests -q

Result: 68 passed
""",
    )
    doc.add_paragraph(
        "Backend tests include strict result LaTeX checks, teaching-step checks, representative step snapshots, and guards against unwanted symbolic or debug artifacts."
    )

    doc.add_heading("7.2 Frontend Tests", level=2)
    add_code(
        doc,
        r"""
cd flutter_app
flutter test

Result: 6 passed, 1 skipped
""",
    )
    doc.add_paragraph(
        "Frontend tests cover responsive breakpoints, home page rendering, example navigation, common screen sizes without layout exceptions, local scroll containers for long formulas and notes, and FT Steps long-formula display on narrow screens. "
        "The optional frontend-backend e2e test is skipped unless explicitly enabled so that frontend CI is not blocked by backend availability."
    )

    doc.add_heading("7.3 Validation Summary", level=2)
    add_table(
        doc,
        ["Validation Area", "Current Evidence"],
        [
            ["Mathematical results", "Backend regression tests verify implemented transform pairs and result_latex output."],
            ["Teaching steps", "Step snapshot tests check representative derivation fragments."],
            ["Output style", "Tests reject Piecewise, RootSum, arg, polar_lift, meijerg, matcher/debug text, and related artifacts."],
            ["Frontend layout", "Widget tests check narrow screens and local scroll containers."],
            ["Integration", "Optional frontend-backend e2e test checks keypad input, backend request, and rendered result."],
        ],
        [2600, 6760],
    )

    doc.add_heading("Chapter 8: Discussion", level=1)
    doc.add_heading("8.1 General User Input Capability", level=2)
    doc.add_paragraph(
        "The app accepts user-entered expressions and can find symbolic Fourier transforms within a defined engineering range. "
        "It is not limited to fixed examples; the examples are used as guided entry points into supported transform families. "
        "For unsupported expressions, the system should return an integral form or error rather than pretending to have a reliable closed form."
    )

    doc.add_heading("8.2 Educational Value", level=2)
    doc.add_paragraph(
        "The system combines readable symbolic output, derivation steps, examples, and numerical visualisation. "
        "This supports both formula recognition and conceptual understanding of how frequency-domain behavior emerges from time-domain signals."
    )

    doc.add_heading("8.3 Engineering Trade-Offs", level=2)
    for item in [
        "Rule-first design improves readability but requires deliberate rule expansion.",
        "A single-file backend is acceptable for prototype iteration but should be modularized later.",
        "Flutter provides portability but requires careful layout handling for math display.",
        "Frontend and backend deployment separation improves operational flexibility but requires configurable API endpoints.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Chapter 9: Conclusion and Future Work", level=1)
    doc.add_heading("9.1 Conclusion", level=2)
    doc.add_paragraph(
        "This project demonstrates a rule-first, distribution-aware Fourier transform engine for engineering education. "
        "It produces readable symbolic results, generates teaching-oriented derivation steps, and integrates them with a Flutter interface that supports examples, LaTeX rendering, and numerical visualisation. "
        "The current implementation is validated by backend and frontend automated tests and includes recent refinements for notation consistency and small-screen formula display."
    )
    doc.add_heading("9.2 Limitations", level=2)
    for item in [
        "The backend covers a defined rule range rather than all possible Fourier transforms.",
        "Complex unsupported expressions may return an integral form or error.",
        "The numerical visualisation is approximate and should be treated as an educational aid.",
        "Current frontend tests detect layout exceptions but are not full screenshot-based visual regression tests.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("9.3 Future Work", level=2)
    for item in [
        "Modularize the backend into rules, formatters, steps, API, and tests packages.",
        "Expand transform-pair coverage and document active versus pending cases.",
        "Add more frontend-backend integration tests.",
        "Add screenshot or golden tests for important layouts.",
        "Improve parser error messages and user-facing unsupported-case explanations.",
        "Support exportable derivation reports for classroom use.",
    ]:
        doc.add_paragraph(item, style="List Bullet")

    doc.add_heading("Appendix A: Existing Report Mapping", level=1)
    doc.add_paragraph(
        "A detailed mapping from the previous project report chapters to this dissertation structure is maintained in:"
    )
    add_code(doc, "docs/word_to_dissertation_mapping.md")

    doc.add_heading("Appendix B: Key Reference Files", level=1)
    add_table(
        doc,
        ["Resource", "Path"],
        [
            ["Backend main file", "backend/backend.py"],
            ["Flutter home page", "flutter_app/lib/main.dart"],
            ["Symbolic result page", "flutter_app/lib/fft/symbol.dart"],
            ["FT Steps page", "flutter_app/lib/fft/step.dart"],
            ["Scrollable formula/text widgets", "flutter_app/lib/scrollable_content.dart"],
            ["Detailed Q&A", "docs/QA.md"],
            ["Brief email Q&A", "docs/QA_email_brief.md"],
            ["Test data reference", "test.md"],
        ],
        [3000, 6360],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
