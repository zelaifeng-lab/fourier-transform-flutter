from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SRC = ROOT / "docs" / "Development_of_Flutter_Apps_for_Fourier_Transform_Dissertation_Draft_CN.md"
OUT = ROOT / "docs" / "Development_of_Flutter_Apps_for_Fourier_Transform_Dissertation_Draft_CN.docx"


def setup_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.18

    for name, size, color in [
        ("Heading 1", 16, RGBColor(46, 116, 181)),
        ("Heading 2", 13, RGBColor(46, 116, 181)),
        ("Heading 3", 12, RGBColor(31, 77, 120)),
    ]:
        style = doc.styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(10)
        style.paragraph_format.space_after = Pt(6)

    code = doc.styles.add_style("CodeBlockCN", 1)
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.space_after = Pt(0)


def add_table_from_buffer(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    table = doc.add_table(rows=1, cols=len(rows[0]))
    table.style = "Table Grid"
    for i, cell_text in enumerate(rows[0]):
        table.rows[0].cells[i].text = cell_text
        for run in table.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
        tc_pr = table.rows[0].cells[i]._tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E8EEF5")
        tc_pr.append(shd)
    for row_data in rows[1:]:
        row = table.add_row()
        for i, value in enumerate(row_data):
            row.cells[i].text = value
    doc.add_paragraph()


def build() -> None:
    doc = Document()
    setup_styles(doc)

    in_code = False
    code_lang = ""
    table_rows: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_rows
        if table_rows:
            cleaned = [row for row in table_rows if not all(set(c.strip()) <= {"-"} for c in row)]
            add_table_from_buffer(doc, cleaned)
            table_rows = []

    for raw in SRC.read_text(encoding="utf-8").splitlines():
        line = raw.rstrip()
        if line.startswith("```"):
            flush_table()
            in_code = not in_code
            code_lang = line.strip("`").strip()
            continue
        if in_code:
            p = doc.add_paragraph(style="CodeBlockCN")
            p.add_run(line)
            continue

        if line.startswith("|") and line.endswith("|"):
            cells = [cell.strip() for cell in line.strip("|").split("|")]
            table_rows.append(cells)
            continue
        flush_table()

        if not line.strip():
            continue
        if line.startswith("# "):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(line[2:].strip())
            r.bold = True
            r.font.size = Pt(22)
            r.font.color.rgb = RGBColor(11, 37, 69)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("- "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        if line[0:3].replace(".", "").isdigit() and ". " in line[:4]:
            doc.add_paragraph(line.split(". ", 1)[1], style="List Number")
            continue

        text = line.replace("**", "")
        doc.add_paragraph(text)

    flush_table()
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()

